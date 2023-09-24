# docx для обработки текстовых файлов .docx
# conda install docx
import docx
import os

def folder_check(path, folder):
    if folder in os.listdir(path):
        pass
    else:
        os.mkdir(f'{path}/{folder}')

def change_font(dirs):
    for dir in dirs:
        doc = docx.Document(dir)

        filename, file_extension = os.path.splitext(os.path.basename(dir))

        for paragraph in doc.paragraphs:
            paragraph_format = paragraph.paragraph_format
            paragraph_format.line_spacing = 1.5

            for run in paragraph.runs:
                font = run.font
                run.font.size = docx.shared.Pt(14)
                run.font.name = 'Times New Roman'

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph_format = paragraph.paragraph_format
                        paragraph_format.line_spacing = 1.5

                        for run in paragraph.runs:
                            font = run.font
                            run.font.size = docx.shared.Pt(14)
                            run.font.name = 'Times New Roman'

        doc.save(f'{os.getcwd()}/Итог/{filename}{file_extension}')

paths = []
folder = os.getcwd()
os.chdir('Исходники')

folder_check(f'{folder}/Исходники', 'Итог')

for root, dirs, files in os.walk(folder):
    for file in files:
        if file.endswith('docx') and not file.startswith('~'):
            paths.append(os.path.join(root, file))

change_font(paths)